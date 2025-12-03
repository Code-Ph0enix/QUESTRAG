import os
from pathlib import Path
from typing import Optional, Dict, Any
from fastapi import UploadFile, HTTPException
import pytesseract
from PIL import Image
import PyPDF2
from docx import Document
from io import BytesIO

from app.utils.file_utils import (
    validate_file_type, validate_file_size, generate_unique_filename, 
    save_upload_file, ALLOWED_IMAGE_TYPES, 
    # ALLOWED_DOC_TYPES, 
    ALLOWED_AUDIO_TYPES
)
from app.config import settings


class FileService:
    """File processing service for images, PDFs, documents, and audio"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        print("✅ FileService initialized")
    
    async def process_image(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Upload image + OCR extraction.
        
        Args:
            file: Uploaded image file
            user_id: User ID (for file organization)
        
        Returns:
            Dict with file_id, path, extracted_text, size
        """
        if not validate_file_type(file, ALLOWED_IMAGE_TYPES):
            raise HTTPException(400, "Invalid image type. Allowed: JPG, PNG, WEBP")
        if not validate_file_size(file):
            raise HTTPException(400, "File too large (max 10MB)")
        
        # Save file
        filename = generate_unique_filename(file.filename)
        filepath = self.upload_dir / "images" / user_id / filename
        await save_upload_file(file, filepath)
        
        # OCR extraction
        try:
            image = Image.open(filepath)
            text = pytesseract.image_to_string(image)
        except Exception as e:
            print(f"⚠️ OCR failed: {e}")
            text = ""
        
        return {
            "file_id": filename,
            "file_path": str(filepath.relative_to(self.upload_dir)),
            "file_type": "image",
            "extracted_text": text.strip(),
            "size": filepath.stat().st_size,
            "original_filename": file.filename
        }
    
    async def process_pdf(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Upload PDF + text extraction.
        
        Args:
            file: Uploaded PDF file
            user_id: User ID
        
        Returns:
            Dict with file_id, path, extracted_text, pages, size
        """
        if file.content_type != "application/pdf":
            raise HTTPException(400, "Invalid PDF file")
        if not validate_file_size(file):
            raise HTTPException(400, "File too large (max 10MB)")
        
        # Save
        filename = generate_unique_filename(file.filename)
        filepath = self.upload_dir / "documents" / user_id / filename
        await save_upload_file(file, filepath)
        
        # Extract text
        text = ""
        pages = 0
        try:
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                pages = len(pdf_reader.pages)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"⚠️ PDF extraction failed: {e}")
        
        return {
            "file_id": filename,
            "file_path": str(filepath.relative_to(self.upload_dir)),
            "file_type": "pdf",
            "extracted_text": text.strip(),
            "pages": pages,
            "size": filepath.stat().st_size,
            "original_filename": file.filename
        }
    
    async def process_docx(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Upload DOCX + text extraction.
        
        Args:
            file: Uploaded DOCX file
            user_id: User ID
        
        Returns:
            Dict with file_id, path, extracted_text, size
        """
        if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise HTTPException(400, "Invalid DOCX file")
        if not validate_file_size(file):
            raise HTTPException(400, "File too large (max 10MB)")
        
        # Save
        filename = generate_unique_filename(file.filename)
        filepath = self.upload_dir / "documents" / user_id / filename
        await save_upload_file(file, filepath)
        
        # Extract
        text = ""
        try:
            # doc = docx.Document(filepath)
            # text = "\n".join([para.text for para in doc.paragraphs])
            doc = Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs])

        except Exception as e:
            print(f"⚠️ DOCX extraction failed: {e}")
        
        return {
            "file_id": filename,
            "file_path": str(filepath.relative_to(self.upload_dir)),
            "file_type": "docx",
            "extracted_text": text.strip(),
            "size": filepath.stat().st_size,
            "original_filename": file.filename
        }
    
    async def process_text_file(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Upload TXT file.
        
        Args:
            file: Uploaded text file
            user_id: User ID
        
        Returns:
            Dict with file_id, path, extracted_text, size
        """
        if file.content_type != "text/plain":
            raise HTTPException(400, "Invalid text file")
        if not validate_file_size(file):
            raise HTTPException(400, "File too large (max 10MB)")
        
        filename = generate_unique_filename(file.filename)
        filepath = self.upload_dir / "documents" / user_id / filename
        await save_upload_file(file, filepath)
        
        text = ""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            print(f"⚠️ Text file read failed: {e}")
        
        return {
            "file_id": filename,
            "file_path": str(filepath.relative_to(self.upload_dir)),
            "file_type": "text",
            "extracted_text": text.strip(),
            "size": filepath.stat().st_size,
            "original_filename": file.filename
        }
    
    # ============================================================================
    # NEW METHOD: Using HuggingFace Transformers Whisper (FREE!)
    # ============================================================================

    async def transcribe_audio(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        """
        Speech-to-text using HuggingFace Transformers Whisper (FREE!).
    
        Args:
            file: Uploaded audio file
            user_id: User ID
    
        Returns:
            Dict with file_id, path, transcription, size
        """
        if not validate_file_type(file, ALLOWED_AUDIO_TYPES):
            raise HTTPException(400, "Invalid audio type. Allowed: MP3, WAV, WEBM, OGG, M4A")
        if not validate_file_size(file):
            raise HTTPException(400, "File too large (max 10MB)")
    
        # Save audio
        filename = generate_unique_filename(file.filename)
        filepath = self.upload_dir / "audio" / user_id / filename
        await save_upload_file(file, filepath)
    
        # Transcribe using HuggingFace Transformers Whisper (FREE!)
        transcription = ""
        try:
            from transformers import pipeline
            import torch
        
            # Lazy load model (only first time)
            if not hasattr(self, '_whisper_pipe'):
                print("🎤 Loading Whisper model (one-time)...")
                device = 0 if torch.cuda.is_available() else -1
                self._whisper_pipe = pipeline(
                    "automatic-speech-recognition",
                    model="openai/whisper-small",  # Small = fast, good accuracy
                    device=device
                )
                print("✅ Whisper model loaded")
        
            # Transcribe
            result = self._whisper_pipe(str(filepath))
            transcription = result["text"]
        
        except Exception as e:
            print(f"⚠️ Whisper transcription failed: {e}")
            raise HTTPException(500, f"Transcription failed: {str(e)}")
    
        return {
            "file_id": filename,
            "file_path": str(filepath.relative_to(self.upload_dir)),
            "file_type": "audio",
            "transcription": transcription,
            "size": filepath.stat().st_size,
            "original_filename": file.filename
        }

    # ============================================================================
    # Old method: OpenAI Whisper API (paid) kept for reference
    # ============================================================================
        # async def transcribe_audio(self, file: UploadFile, user_id: str) -> Dict[str, Any]:
        #     """
        #     Speech-to-text using OpenAI Whisper API.
        
        #     Args:
        #         file: Uploaded audio file
        #         user_id: User ID
        
        #     Returns:
        #         Dict with file_id, path, transcription, size
        #     """
        #     if not validate_file_type(file, ALLOWED_AUDIO_TYPES):
        #         raise HTTPException(400, "Invalid audio type. Allowed: MP3, WAV, WEBM, OGG, M4A")
        #     if not validate_file_size(file):
        #         raise HTTPException(400, "File too large (max 10MB)")

        #     # Save audio
        #     filename = generate_unique_filename(file.filename)
        #     filepath = self.upload_dir / "audio" / user_id / filename
        #     await save_upload_file(file, filepath)

        #     # Transcribe using OpenAI Whisper API
        #     transcription = ""
        #     try:
        #         from openai import OpenAI
        #         client = OpenAI(api_key=settings.OPENAI_API_KEY)

        #         with open(filepath, "rb") as audio_file:
        #             transcript = client.audio.transcriptions.create(
        #                 model="whisper-1",
        #                 file=audio_file,
        #                 language="en"  # Change if needed
        #             )

        #         transcription = transcript.text
        #     except Exception as e:
        #         print(f"⚠️ Whisper transcription failed: {e}")
        #         raise HTTPException(500, f"Transcription failed: {str(e)}")

        #     return {
        #         "file_id": filename,
        #         "file_path": str(filepath.relative_to(self.upload_dir)),
        #         "file_type": "audio",
        #         "transcription": transcription,
        #         "size": filepath.stat().st_size,
        #         "original_filename": file.filename
        #     }
    
    def delete_file(self, file_path: str, user_id: str) -> bool:
        """
        Delete uploaded file.
        
        Args:
            file_path: Relative file path (from upload_dir)
            user_id: User ID (for security check)
        
        Returns:
            bool: True if deleted
        """
        try:
            # Security: Ensure file belongs to user
            if user_id not in file_path:
                return False
            
            full_path = self.upload_dir / file_path
            if full_path.exists() and full_path.is_file():
                full_path.unlink()
                return True
            return False
        except Exception as e:
            print(f"⚠️ File deletion failed: {e}")
            return False


# ============================================================================
# GLOBAL SERVICE INSTANCE
# ============================================================================

file_service = FileService()